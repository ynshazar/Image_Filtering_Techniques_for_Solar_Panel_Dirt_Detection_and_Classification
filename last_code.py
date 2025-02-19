import os
import subprocess
import sys

import numpy as np
import cv2
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.metrics import classification_report, confusion_matrix, roc_curve, auc, precision_recall_curve
from tensorflow.keras.utils import to_categorical
from tensorflow.keras import layers, models
from sklearn.model_selection import train_test_split
from docx import Document
from docx.shared import Inches

# Görsel boyutları (100x100)
image_size = (100, 100)

# Filtre uygulama fonksiyonları
def apply_filters(image):
    filters = {}
    filters['original'] = image
    filters['gaussian_blur'] = cv2.GaussianBlur(image, (5, 5), 0)
    filters['unsharp_mask'] = cv2.addWeighted(image, 1.5, cv2.GaussianBlur(image, (5, 5), 10), -0.5, 0)
    filters['bilateral_filtered'] = cv2.bilateralFilter(image, 9, 75, 75)
    filters['median_filtered'] = cv2.medianBlur(image, 5)
    return filters

# Verileri yükleme ve filtreleme fonksiyonu
def load_and_filter_images(directory):
    images = []
    labels = []
    filter_types = []

    for label, folder in enumerate(['clean', 'dirty']):
        folder_path = os.path.join(directory, folder)
        for filename in os.listdir(folder_path):
            if filename.endswith('.png') or filename.endswith('.jpg'):
                img_path = os.path.join(folder_path, filename)
                img = cv2.imread(img_path, cv2.IMREAD_COLOR)
                if img is not None:
                    img = cv2.resize(img, image_size)
                    filtered_images = apply_filters(img)
                    for filter_name, filtered_img in filtered_images.items():
                        images.append(filtered_img)
                        labels.append(label)
                        filter_types.append(filter_name)
    return np.array(images), np.array(labels), filter_types

# Verilerin bulunduğu ana dizin
main_dir = r'images'

# Verileri yükle
X, y, filter_types = load_and_filter_images(main_dir)

# Etiketleri one-hot encoding'e dönüştür
y = to_categorical(y, num_classes=2)

# Eğitim ve test verisi olarak ayır
X_train, X_test, y_train, y_test, filters_train, filters_test = train_test_split(X, y, filter_types, test_size=0.2, random_state=42)

# Modeli oluşturma
input_shape = (image_size[0], image_size[1], 3)

model = models.Sequential([
    layers.Conv2D(32, (3, 3), activation='relu', input_shape=input_shape),
    layers.MaxPooling2D((2, 2)),
    layers.Conv2D(64, (3, 3), activation='relu'),
    layers.MaxPooling2D((2, 2)),
    layers.Conv2D(128, (3, 3), activation='relu'),
    layers.MaxPooling2D((2, 2)),
    layers.Flatten(),
    layers.Dense(128, activation='relu'),
    layers.Dense(2, activation='softmax')
])

# Modeli derle
model.compile(optimizer='adam', loss='categorical_crossentropy', metrics=['accuracy'])

# Model özetini göster
model.summary()

# Modeli eğit
history = model.fit(X_train, y_train, epochs=25, batch_size=32, validation_data=(X_test, y_test))

# Modeli değerlendir
test_loss, test_accuracy = model.evaluate(X_test, y_test)
print(f"Test Loss: {test_loss}")
print(f"Test Accuracy: {test_accuracy}")


# Tahmin yapma
predictions = model.predict(X_test)
y_pred = np.argmax(predictions, axis=1)
y_true = np.argmax(y_test, axis=1)

# Performans metrikleri
classification_rep = classification_report(y_true, y_pred)
conf_matrix = confusion_matrix(y_true, y_pred)

# ROC Eğrisi
fpr, tpr, _ = roc_curve(y_true, predictions[:, 1])
roc_auc = auc(fpr, tpr)

# Precision-Recall Eğrisi
precision, recall, _ = precision_recall_curve(y_true, predictions[:, 1])

# Word belgesi oluşturma
doc = Document()
doc.add_heading('Model Training and Evaluation Results', level=1)

# Performans metriklerini ekle
doc.add_heading('Classification Report:', level=2)
doc.add_paragraph(classification_rep)

doc.add_heading('Confusion Matrix:', level=2)
plt.figure(figsize=(6, 6))
sns.heatmap(conf_matrix, annot=True, fmt='d', cmap='Blues', xticklabels=['Clean', 'Dirty'], yticklabels=['Clean', 'Dirty'])
plt.title('Confusion Matrix')
plt.xlabel('Predicted')
plt.ylabel('True')
plt.savefig('confusion_matrix.png')
plt.close()
doc.add_picture('confusion_matrix.png', width=Inches(5))

# ROC Eğrisi ekle
doc.add_heading('ROC Curve:', level=2)
plt.figure()
plt.plot(fpr, tpr, color='blue', lw=2, label=f'ROC Curve (AUC = {roc_auc:.2f})')
plt.plot([0, 1], [0, 1], color='gray', lw=1, linestyle='--')
plt.title('Receiver Operating Characteristic')
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.legend(loc='lower right')
plt.savefig('roc_curve.png')
plt.close()
doc.add_picture('roc_curve.png', width=Inches(5))

# Precision-Recall Eğrisi ekle
doc.add_heading('Precision-Recall Curve:', level=2)
plt.figure()
plt.plot(recall, precision, color='green', lw=2)
plt.title('Precision-Recall Curve')
plt.xlabel('Recall')
plt.ylabel('Precision')
plt.savefig('precision_recall_curve.png')
plt.close()
doc.add_picture('precision_recall_curve.png', width=Inches(5))

# Eğitim/doğrulama kayıpları ve doğrulukları ekle
doc.add_heading('Training vs Validation Loss:', level=2)
plt.figure()
plt.plot(history.history['loss'], label='Training Loss')
plt.plot(history.history['val_loss'], label='Validation Loss')
plt.title('Training vs Validation Loss')
plt.xlabel('Epoch')
plt.ylabel('Loss')
plt.legend()
plt.savefig('training_loss.png')
plt.close()
doc.add_picture('training_loss.png', width=Inches(5))

doc.add_heading('Training vs Validation Accuracy:', level=2)
plt.figure()
plt.plot(history.history['accuracy'], label='Training Accuracy')
plt.plot(history.history['val_accuracy'], label='Validation Accuracy')
plt.title('Training vs Validation Accuracy')
plt.xlabel('Epoch')
plt.ylabel('Accuracy')
plt.legend()
plt.savefig('training_accuracy.png')
plt.close()
doc.add_picture('training_accuracy.png', width=Inches(5))

# Word belgesini kaydet
doc.save('model_results.docx')

print("All results have been saved to 'model_results.docx'.")
